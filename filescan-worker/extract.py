"""Text extraction for filescan-worker (G16, Phase 3): PDF/DOCX/XLSX/CSV
uploads reduced to a list of (location, text) pairs so each becomes an
independent TextUnit for core/detect's pipeline -- one unit per
page/paragraph/row rather than one giant blob, so a DetectedSpan's
`unit_id` can be traced back to roughly where in the document a
finding lives (see scan.py).
"""
from __future__ import annotations

import csv
import io

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

SUPPORTED_CONTENT_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "text/csv": "csv",
}
SUPPORTED_EXTENSIONS = {"pdf", "docx", "xlsx", "csv"}


class UnsupportedFileTypeError(Exception):
    def __init__(self, content_type: str | None, filename: str | None) -> None:
        super().__init__(f"unsupported file type: content_type={content_type!r} filename={filename!r}")
        self.content_type = content_type
        self.filename = filename


def detect_kind(filename: str | None, content_type: str | None) -> str:
    """Content-type wins when it's one we recognize; otherwise fall
    back to the filename extension (browsers/clients are inconsistent
    about setting a precise multipart Content-Type for office docs)."""
    if content_type in SUPPORTED_CONTENT_TYPES:
        return SUPPORTED_CONTENT_TYPES[content_type]
    ext = filename.rsplit(".", 1)[-1].lower() if filename and "." in filename else ""
    if ext in SUPPORTED_EXTENSIONS:
        return ext
    raise UnsupportedFileTypeError(content_type, filename)


def extract_units(data: bytes, kind: str) -> list[tuple[str, str]]:
    """Returns [(location, text), ...]. `location` is a human-readable
    pointer (page/paragraph/row index) for the scan report -- never a
    sensitive value itself."""
    if kind == "pdf":
        return _extract_pdf(data)
    if kind == "docx":
        return _extract_docx(data)
    if kind == "xlsx":
        return _extract_xlsx(data)
    if kind == "csv":
        return _extract_csv(data)
    raise UnsupportedFileTypeError(None, None)


def _extract_pdf(data: bytes) -> list[tuple[str, str]]:
    reader = PdfReader(io.BytesIO(data))
    units: list[tuple[str, str]] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            units.append((f"page[{i}]", text))
    return units


def _extract_docx(data: bytes) -> list[tuple[str, str]]:
    doc = DocxDocument(io.BytesIO(data))
    units: list[tuple[str, str]] = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            units.append((f"paragraph[{i}]", para.text))
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            row_text = "\t".join(cell.text for cell in row.cells)
            if row_text.strip():
                units.append((f"table[{ti}].row[{ri}]", row_text))
    return units


def _extract_xlsx(data: bytes) -> list[tuple[str, str]]:
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    units: list[tuple[str, str]] = []
    for sheet in wb.worksheets:
        for ri, row in enumerate(sheet.iter_rows(values_only=True)):
            row_text = "\t".join(str(c) for c in row if c is not None)
            if row_text.strip():
                units.append((f"{sheet.title}!row[{ri}]", row_text))
    return units


def _extract_csv(data: bytes) -> list[tuple[str, str]]:
    text = data.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    units: list[tuple[str, str]] = []
    for ri, row in enumerate(reader):
        row_text = "\t".join(row)
        if row_text.strip():
            units.append((f"row[{ri}]", row_text))
    return units
